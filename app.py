import time
import streamlit as st
import json
import re
from docx import Document
from io import BytesIO
from docx.shared import Pt
from datetime import datetime


# ─────────────────────────────────────────────
#  BUSINESS LOGIC  (unchanged)
# ─────────────────────────────────────────────

def evaluate_application(data):
    byggtype = data.get("byggtype") or "Ukjent tiltak"
    areal = data.get("areal") or 0
    avstand = data.get("avstand") or 0
    regulert = data.get("regulert") or "Ikke oppgitt"
    dispensasjon = data.get("dispensasjon") or "Ikke oppgitt"
    eiendomstype = data.get("eiendomstype") or "ukjent eiendomstype"
    mønehøyde = data.get("mønehøyde") or 0
    vinduer_mot_nabo = data.get("vinduer_mot_nabo") or "Ikke oppgitt"
    frittstående = data.get("frittstående") or "Ikke oppgitt"
    nabovarsel_sendt = data.get("nabovarsel_sendt") or "Ikke oppgitt"

    summary = (
        f"Søknaden gjelder {byggtype.lower()} på en {eiendomstype.lower()}. "
        f"Areal: {areal} m². Avstand til nabogrense: {avstand} meter. "
        f"Regulert: {regulert}. Dispensasjon: {dispensasjon}. "
        f"Mønehøyde: {mønehøyde} meter. Vinduer mot nabo: {vinduer_mot_nabo}. "
        f"Frittstående: {frittstående}. "
        f"Nabovarsel sendt: {nabovarsel_sendt}."
    )

    if (
        areal <= 0
        or avstand <= 0
        or mønehøyde <= 0
        or vinduer_mot_nabo not in ["Ja", "Nei"]
        or frittstående not in ["Ja", "Nei"]
        or nabovarsel_sendt not in ["Ja", "Nei"]
    ):
        return {
            "summary": summary,
            "decision": (
                "Forslag: Søknaden mangler nødvendig informasjon og må suppleres.\n\n"
                "Begrunnelse:\n"
                "- Areal, avstand til nabogrense og mønehøyde må være oppgitt med gyldige verdier.\n"
                "- Det må også være opplyst om vinduer mot nabo, om tiltaket er frittstående, og om nabovarsel er sendt.\n\n"
                "Løsning:\n"
                "- Legg inn komplett informasjon før søknaden vurderes videre."
            ),
            "rule_checks": [
                {"label": "Areal er oppgitt", "ok": areal > 0},
                {"label": "Avstand til nabogrense er oppgitt", "ok": avstand > 0},
                {"label": "Mønehøyde er oppgitt", "ok": mønehøyde > 0},
                {"label": "Vinduer mot nabo er oppgitt", "ok": vinduer_mot_nabo in ["Ja", "Nei"]},
                {"label": "Frittstående er oppgitt", "ok": frittstående in ["Ja", "Nei"]},
                {"label": "Nabovarsel er oppgitt", "ok": nabovarsel_sendt in ["Ja", "Nei"]},
            ],
            "indicators": {
                "status": "Krever mer info",
                "struktur": "Søknaden er ufullstendig og må suppleres før videre vurdering.",
                "regelreferanse": "Saken må ha tilstrekkelig beslutningsgrunnlag før kommunen kan behandle den videre.",
                "konsistens": "Systemet kan ikke gi en endelig vurdering når sentrale opplysninger mangler.",
            },
        }

    areal_ok = areal <= 50
    er_garasje = byggtype.strip().lower() == "garasje"
    mønehøyde_ok = mønehøyde <= 5
    vinduer_ok = vinduer_mot_nabo == "Nei"
    frittstående_ok = frittstående == "Ja"
    nabovarsel_ok = nabovarsel_sendt == "Ja"

    if er_garasje and areal_ok and mønehøyde_ok and vinduer_ok and frittstående_ok:
        påkrevd_avstand = 1
    else:
        påkrevd_avstand = 4

    avstand_ok = avstand >= påkrevd_avstand
    antall_avvik = sum([
        not avstand_ok, not areal_ok, not mønehøyde_ok,
        not vinduer_ok, not frittstående_ok, not nabovarsel_ok,
    ])
    søker_dispensasjon = dispensasjon == "Ja"

    if avstand_ok and areal_ok and mønehøyde_ok and vinduer_ok and frittstående_ok and nabovarsel_ok:
        return {
            "summary": summary,
            "decision": (
                "Forslag: Søknaden kan godkjennes.\n\n"
                "Begrunnelse:\n"
                f"- Tiltaket oppfyller krav til avstand (≥ {påkrevd_avstand} meter i denne vurderingen).\n"
                "- Tiltaket er innenfor relevant vurdering av størrelse i prototypen.\n"
                "- Mønehøyden er innenfor forenklet grense på 5 meter.\n"
                "- Tiltaket er oppgitt uten vinduer mot nabo.\n"
                "- Tiltaket er frittstående og nabovarsel er sendt.\n\n"
                "Anbefaling:\n"
                "- Fortsett ordinær behandling og verifiser øvrige krav (reguleringsplan mv.)."
            ),
            "rule_checks": [
                {"label": f"Avstand oppfyller krav på minst {påkrevd_avstand} meter", "ok": avstand_ok},
                {"label": "Tiltaket er innenfor størrelsesgrensen på 50 m²", "ok": areal_ok},
                {"label": "Mønehøyde er maks 5 meter", "ok": mønehøyde_ok},
                {"label": "Tiltaket har ikke vinduer mot nabo", "ok": vinduer_ok},
                {"label": "Tiltaket er frittstående", "ok": frittstående_ok},
                {"label": "Nabovarsel er sendt", "ok": nabovarsel_ok},
            ],
            "indicators": {
                "status": "Godkjent",
                "struktur": "Søknaden fremstår som komplett innenfor de forenklede kriteriene.",
                "regelreferanse": "Tiltaket ser ut til å være innenfor forenklet vurdering av avstand og størrelse. Peker mot pbl. § 29-4.",
                "konsistens": "Konklusjonen er logisk ut fra oppgitt areal og avstand til nabogrense.",
            },
        }

    if (not avstand_ok or not areal_ok or not mønehøyde_ok
            or not vinduer_ok or not frittstående_ok or not nabovarsel_ok) and søker_dispensasjon:
        reasons, videre_behandling = [], []
        if not avstand_ok:
            reasons.append(f"Avstand til nabogrense er mindre enn {påkrevd_avstand} meter")
            videre_behandling.append(f"Vurder dispensasjon fra avstandskravet på {påkrevd_avstand} meter etter pbl. § 19-2")
        if not areal_ok:
            reasons.append("Tiltaket overstiger forenklet størrelsesgrense (50 m²)")
            videre_behandling.append("Vurder om tiltakets størrelse krever endring eller nærmere behandling")
        if not mønehøyde_ok:
            reasons.append("Mønehøyden overstiger 5 meter")
            videre_behandling.append("Vurder om mønehøyden må reduseres eller om det må søkes dispensasjon")
        if not vinduer_ok:
            reasons.append("Tiltaket er oppgitt med vinduer mot nabo")
            videre_behandling.append("Vurder om vinduer mot nabo må fjernes eller endres")
        if not frittstående_ok:
            reasons.append("Tiltaket er ikke oppgitt som frittstående")
            videre_behandling.append("Vurder om tiltaket må omprosjekteres som frittstående")
        if not nabovarsel_ok:
            reasons.append("Nabovarsel er ikke sendt")
            videre_behandling.append("Send nabovarsel før videre behandling av saken")

        return {
            "summary": summary,
            "decision": (
                "Forslag: Søknaden krever dispensasjon.\n\n"
                "Begrunnelse:\n- " + "\n- ".join(reasons)
                + "\n\nLøsning/videre behandling:\n- " + "\n- ".join(videre_behandling)
                + "\n- Innhent nabovarsel/uttalelser og vurder ulemper og fordeler."
            ),
            "rule_checks": [
                {"label": f"Avstand oppfyller krav på minst {påkrevd_avstand} meter", "ok": avstand_ok},
                {"label": "Tiltaket er innenfor størrelsesgrensen på 50 m²", "ok": areal_ok},
                {"label": "Mønehøyde er maks 5 meter", "ok": mønehøyde_ok},
                {"label": "Tiltaket har ikke vinduer mot nabo", "ok": vinduer_ok},
                {"label": "Tiltaket er frittstående", "ok": frittstående_ok},
                {"label": "Nabovarsel er sendt", "ok": nabovarsel_ok},
            ],
            "indicators": {
                "status": "Krever dispensasjon",
                "antall_avvik": f"Søknaden har {antall_avvik} registrerte avvik som må vurderes nærmere.",
                "struktur": "Søknaden inneholder avvik fra de forenklede hovedkravene og må vurderes nærmere.",
                "regelreferanse": "Peker mot pbl. § 19-2 om dispensasjon, og kan berøre § 29-4 om avstandskrav.",
                "konsistens": "Alle registrerte avvik tas med i videre dispensasjonsvurdering.",
            },
        }

    reasons, suggestions = [], []
    if not avstand_ok:
        reasons.append(f"Avstand til nabogrense er mindre enn {påkrevd_avstand} meter")
        suggestions.append(f"Øk avstanden til minst {påkrevd_avstand} meter fra nabogrense")
    if not areal_ok:
        reasons.append("Tiltaket overstiger forenklet størrelsesgrense (50 m²)")
        suggestions.append("Reduser arealet til maks 50 m²")
    if not mønehøyde_ok:
        reasons.append("Mønehøyden overstiger 5 meter")
        suggestions.append("Reduser mønehøyden til maks 5 meter")
    if not vinduer_ok:
        reasons.append("Tiltaket er oppgitt med vinduer mot nabo")
        suggestions.append("Fjern eller endre vinduer som vender mot nabo")
    if not frittstående_ok:
        reasons.append("Tiltaket er ikke oppgitt som frittstående")
        suggestions.append("Omprosjekter tiltaket som frittstående")
    if not nabovarsel_ok:
        reasons.append("Nabovarsel er ikke sendt")
        suggestions.append("Send nabovarsel før saken behandles videre")
    if not suggestions:
        suggestions.append("Vurder å søke dispensasjon dersom vilkårene kan oppfylles")

    return {
        "summary": summary,
        "decision": (
            "Forslag: Søknaden bør avslås.\n\n"
            "Begrunnelse:\n- " + "\n- ".join(reasons)
            + "\n\nMulige løsninger:\n- " + "\n- ".join(suggestions)
            + "\n- Alternativt kan søker vurdere om det finnes grunnlag for dispensasjon."
        ),
        "rule_checks": [
            {"label": f"Avstand oppfyller krav på minst {påkrevd_avstand} meter", "ok": avstand_ok},
            {"label": "Tiltaket er innenfor størrelsesgrensen på 50 m²", "ok": areal_ok},
            {"label": "Mønehøyde er maks 5 meter", "ok": mønehøyde_ok},
            {"label": "Tiltaket har ikke vinduer mot nabo", "ok": vinduer_ok},
            {"label": "Tiltaket er frittstående", "ok": frittstående_ok},
            {"label": "Nabovarsel er sendt", "ok": nabovarsel_ok},
        ],
        "indicators": {
            "status": "Avslag",
            "antall_avvik": f"Søknaden har {antall_avvik} registrerte avvik fra hovedkravene.",
            "struktur": "Søknaden bryter ett eller flere av de forenklede hovedkravene i prototypen.",
            "regelreferanse": "Peker mot pbl. § 29-4 om plassering og avstand.",
            "konsistens": "Systemet gir avslag fordi alle oppgitte avvik tas med i vurderingen.",
        },
    }


def parse_uploaded_application(uploaded_file):
    if uploaded_file is None:
        return None
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".docx"):
        doc = Document(uploaded_file)
        raw_text = "\n".join([p.text for p in doc.paragraphs])
    else:
        raw_text = uploaded_file.read().decode("utf-8")
    if file_name.endswith(".json"):
        parsed = json.loads(raw_text)
        return {
            "byggtype": parsed.get("byggtype", "Garasje"),
            "areal": float(parsed.get("areal", 0)),
            "avstand": float(parsed.get("avstand", 0)),
            "regulert": parsed.get("regulert", "Ja"),
            "dispensasjon": parsed.get("dispensasjon", "Nei"),
            "eiendomstype": parsed.get("eiendomstype", "Enebolig"),
            "mønehøyde": float(parsed.get("mønehøyde", 0)),
            "vinduer_mot_nabo": parsed.get("vinduer_mot_nabo", "Nei"),
            "frittstående": parsed.get("frittstående", "Ja"),
            "nabovarsel_sendt": parsed.get("nabovarsel_sendt", "Ja"),
        }

    def find_value(labels, default=""):
        if isinstance(labels, str):
            labels = [labels]
        for label in labels:
            pattern = rf"{label}\s*:\s*(.+)"
            match = re.search(pattern, raw_text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return default

    tiltak_text = find_value(["byggtype", "tiltak"], "Garasje")
    størrelse_text = find_value(["areal", "størrelse"], "0")
    avstand_text = find_value(["avstand", "avstand til nabogrense"], "0")
    regulering_text = find_value(["regulert", "reguleringsstatus"], "Ja")
    dispensasjon_text = find_value(["dispensasjon", "dispensasjon søkt"], "Nei")
    eiendomstype_text = find_value(["eiendomstype"], "Enebolig")
    mønehøyde_text = find_value(["mønehøyde"], "0")
    vinduer_text = find_value(["vinduer mot nabo"], "Nei")
    frittstående_text = find_value(["frittstående", "er tiltaket frittstående"], "Ja")
    nabovarsel_text = find_value(["nabovarsel sendt", "er nabovarsel sendt"], "Ja")

    tiltak_lower = tiltak_text.lower()
    if "garasje" in tiltak_lower:
        byggtype = "Garasje"
    elif "tilbygg" in tiltak_lower:
        byggtype = "Tilbygg"
    else:
        byggtype = tiltak_text.title()

    størrelse_match = re.search(r"(\d+[\.,]?\d*)", størrelse_text)
    avstand_match = re.search(r"(\d+[\.,]?\d*)", avstand_text)
    mønehøyde_match = re.search(r"(\d+[\.,]?\d*)", mønehøyde_text)

    areal = float(størrelse_match.group(1).replace(",", ".")) if størrelse_match else 0.0
    avstand = float(avstand_match.group(1).replace(",", ".")) if avstand_match else 0.0
    mønehøyde = float(mønehøyde_match.group(1).replace(",", ".")) if mønehøyde_match else 0.0

    regulering_lower = regulering_text.lower()
    regulert = regulering_text.title() if regulering_lower in ["ja", "nei"] else "Ja"
    dispensasjon_lower = dispensasjon_text.lower()
    dispensasjon = dispensasjon_text.title() if dispensasjon_lower in ["ja", "nei"] else ("Ja" if "ja" in dispensasjon_lower else "Nei")

    return {
        "byggtype": byggtype, "areal": areal, "avstand": avstand,
        "regulert": regulert, "dispensasjon": dispensasjon,
        "eiendomstype": eiendomstype_text.title(), "mønehøyde": mønehøyde,
        "vinduer_mot_nabo": vinduer_text.title(), "frittstående": frittstående_text.title(),
        "nabovarsel_sendt": nabovarsel_text.title(),
    }


def build_nabovarsel_template(data):
    doc = Document()
    title = doc.add_paragraph()
    title_run = title.add_run("NABOVARSEL")
    title_run.bold = True
    title_run.font.size = Pt(16)
    doc.add_paragraph("Mal generert fra byggesaksprototypen.")
    doc.add_paragraph("")

    p = doc.add_paragraph(); p.add_run("1. Tiltak på eiendommen").bold = True
    t1 = doc.add_table(rows=5, cols=2); t1.style = "Table Grid"
    t1.cell(0, 0).text = "Tiltakstype";           t1.cell(0, 1).text = str(data.get("byggtype", ""))
    t1.cell(1, 0).text = "Eiendomstype";          t1.cell(1, 1).text = str(data.get("eiendomstype", ""))
    t1.cell(2, 0).text = "Areal";                 t1.cell(2, 1).text = f"{data.get('areal', '')} m²"
    t1.cell(3, 0).text = "Avstand til nabogrense"; t1.cell(3, 1).text = f"{data.get('avstand', '')} meter"
    t1.cell(4, 0).text = "Mønehøyde";             t1.cell(4, 1).text = f"{data.get('mønehøyde', '')} meter"
    doc.add_paragraph("")

    p = doc.add_paragraph(); p.add_run("2. Opplysninger om tiltaket").bold = True
    t2 = doc.add_table(rows=4, cols=2); t2.style = "Table Grid"
    t2.cell(0, 0).text = "Vinduer mot nabo";      t2.cell(0, 1).text = str(data.get("vinduer_mot_nabo", ""))
    t2.cell(1, 0).text = "Frittstående";           t2.cell(1, 1).text = str(data.get("frittstående", ""))
    t2.cell(2, 0).text = "Dispensasjon søkt";     t2.cell(2, 1).text = str(data.get("dispensasjon", ""))
    t2.cell(3, 0).text = "Nabovarsel sendt";       t2.cell(3, 1).text = str(data.get("nabovarsel_sendt", ""))
    doc.add_paragraph("")

    p = doc.add_paragraph(); p.add_run("3. Beskrivelse").bold = True
    doc.add_paragraph(
        f"Det varsles om planlagt {str(data.get('byggtype','')).lower()} på {str(data.get('eiendomstype','')).lower()}. "
        f"Areal: {data.get('areal','')} m², avstand til nabogrense: {data.get('avstand','')} meter, mønehøyde: {data.get('mønehøyde','')} meter."
    )
    doc.add_paragraph("Naboer bes sende eventuelle merknader innen gjeldende frist.")
    doc.add_paragraph("")

    p = doc.add_paragraph(); p.add_run("4. Mottaker og merknader").bold = True
    t3 = doc.add_table(rows=4, cols=2); t3.style = "Table Grid"
    for i, (k, v) in enumerate([("Til nabo/gjenboer","___________"),("Adresse","___________"),("Frist for merknader","Innen 2 uker"),("Merknader sendes til","___________")]):
        t3.cell(i, 0).text = k; t3.cell(i, 1).text = v
    doc.add_paragraph("")

    p = doc.add_paragraph(); p.add_run("5. Underskrift").bold = True
    t4 = doc.add_table(rows=3, cols=2); t4.style = "Table Grid"
    for i, (k, v) in enumerate([("Sted","___________"),("Dato","___________"),("Underskrift","___________")]):
        t4.cell(i, 0).text = k; t4.cell(i, 1).text = v

    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer


def build_summary_docx(data, result):
    doc = Document()
    title = doc.add_paragraph()
    title_run = title.add_run("Oppsummering av byggesak")
    title_run.bold = True; title_run.font.size = Pt(16)
    doc.add_paragraph(f"Generert: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("")

    for heading, content in [
        ("Sammendrag", result.get("summary", "")),
        ("Kategori", result.get("indicators", {}).get("status", "Ukjent")),
        ("Forslag til vedtak", result.get("decision", "")),
    ]:
        p = doc.add_paragraph(); p.add_run(heading).bold = True
        doc.add_paragraph(content)

    p = doc.add_paragraph(); p.add_run("Indikatorer").bold = True
    ind = result.get("indicators", {})
    doc.add_paragraph(f"Status: {ind.get('status','Ukjent')}")
    if ind.get("antall_avvik"):
        doc.add_paragraph(f"Omfang av avvik: {ind.get('antall_avvik')}")
    doc.add_paragraph(f"Struktur: {ind.get('struktur','')}")
    doc.add_paragraph(f"Regelreferanse: {ind.get('regelreferanse','')}")
    doc.add_paragraph(f"Konsistens: {ind.get('konsistens','')}")

    p = doc.add_paragraph(); p.add_run("Regelkontroll").bold = True
    for rule in result.get("rule_checks", []):
        icon = "OK" if rule.get("ok") else "AVVIK"
        doc.add_paragraph(f"[{icon}] {rule.get('label')}")

    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer


# ─────────────────────────────────────────────
#  PAGE CONFIG & GLOBAL CSS
# ─────────────────────────────────────────────

st.set_page_config(
    page_title="Byggesaksprototype",
    layout="wide",
    page_icon="🏗️",
    initial_sidebar_state="expanded",
)

st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;1,9..40,300&display=swap" rel="stylesheet">

<style>
/* ── Root tokens ─────────────────────────────────────── */
:root {
  --navy:   #0d1b2a;
  --navy2:  #162032;
  --navy3:  #1e2d3d;
  --teal:   #00c9a7;
  --teal2:  #00a88d;
  --amber:  #f59e0b;
  --red:    #ef4444;
  --green:  #10b981;
  --blue:   #3b82f6;
  --surface:#0d1b2a;
  --border: rgba(138,175,200,0.35);
  --text:   #dbeafe;
  --muted:  #8aafc8;
  --white:  #ffffff;
  --radius: 12px;
  --shadow: 0 4px 24px rgba(13,27,42,0.10);
  --shadow-lg: 0 8px 40px rgba(13,27,42,0.16);
}

/* ── Global reset ────────────────────────────────────── */
html, body, [class*="css"] {
  font-family: 'DM Sans', sans-serif !important;
  color: var(--text) !important;
  background: var(--surface) !important;
}

.stApp {
  background: linear-gradient(135deg, #0d1b2a 0%, #111827 55%, #162032 100%) !important;
  color: var(--text) !important;
}

.main,
.block-container,
section.main,
[data-testid="stAppViewContainer"] {
  background: transparent !important;
  color: var(--text) !important;
}

.stMarkdown,
.stMarkdown p,
.stMarkdown span,
.stMarkdown div {
  color: inherit;
}

/* ── Sidebar ─────────────────────────────────────────── */
[data-testid="stSidebar"] {
  background: var(--navy) !important;
  border-right: none !important;
}
[data-testid="stSidebar"] * {
  color: #c8d8e8 !important;
  font-family: 'DM Sans', sans-serif !important;
}
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] strong {
  color: var(--white) !important;
  font-family: 'Syne', sans-serif !important;
}
[data-testid="stSidebar"] .stMarkdown p {
  font-size: 0.88rem !important;
  line-height: 1.65 !important;
  color: #9ab2c8 !important;
}
[data-testid="stSidebar"] hr {
  border-color: #2a3f55 !important;
}

/* ── Hide default Streamlit chrome ──────────────────── */
#MainMenu, footer, header { visibility: hidden !important; }
[data-testid="stToolbar"] { display: none !important; }
[data-testid="stDecoration"] { display: none !important; }
.block-container { padding-top: 1.5rem !important; max-width: 1160px !important; }

/* Keep sidebar visible and remove collapse controls */
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapseButton"] {
  display: none !important;
}

[data-testid="stSidebar"] {
  min-width: 18rem !important;
  max-width: 18rem !important;
}

/* ── Typography ──────────────────────────────────────── */
h1, h2, h3, h4 {
  font-family: 'Syne', sans-serif !important;
  letter-spacing: -0.02em !important;
}

/* ── Dividers ────────────────────────────────────────── */
hr { border-color: var(--border) !important; margin: 1.6rem 0 !important; }

/* ── Selectbox & inputs ──────────────────────────────── */
[data-testid="stSelectbox"] label,
[data-testid="stNumberInput"] label,
[data-testid="stRadio"] label,
[data-testid="stFileUploader"] label {
  font-size: 0.82rem !important;
  font-weight: 700 !important;
  text-transform: uppercase !important;
  letter-spacing: 0.07em !important;
  color: #8aafc8 !important;
}

[data-testid="stSelectbox"] label *,
[data-testid="stNumberInput"] label *,
[data-testid="stRadio"] label *,
[data-testid="stFileUploader"] label * {
  color: #8aafc8 !important;
}
 [data-testid="stSelectbox"] > div > div,
 [data-testid="stNumberInput"] input {
   border: 1.5px solid var(--border) !important;
   border-radius: 8px !important;
   background: var(--white) !important;
   color: #1a2535 !important;
   font-size: 0.95rem !important;
   transition: border-color 0.2s !important;
 }

 [data-testid="stSelectbox"] * {
   color: #1a2535 !important;
 }

 [data-testid="stNumberInput"] input {
   color: #1a2535 !important;
   -webkit-text-fill-color: #1a2535 !important;
 }

 [data-baseweb="select"] span {
   color: #1a2535 !important;
 }
[data-testid="stSelectbox"] > div > div:hover,
[data-testid="stNumberInput"] input:focus {
  border-color: var(--teal) !important;
  box-shadow: 0 0 0 3px rgba(0,201,167,0.12) !important;
}

/* ── Buttons ─────────────────────────────────────────── */
[data-testid="stButton"] > button {
  background: var(--teal) !important;
  color: var(--navy) !important;
  font-family: 'Syne', sans-serif !important;
  font-weight: 700 !important;
  font-size: 0.95rem !important;
  letter-spacing: 0.04em !important;
  border: none !important;
  border-radius: 8px !important;
  padding: 0.65rem 2.2rem !important;
  transition: background 0.2s, transform 0.15s, box-shadow 0.2s !important;
  box-shadow: 0 2px 12px rgba(0,201,167,0.30) !important;
}
[data-testid="stButton"] > button:hover {
  background: var(--teal2) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 6px 20px rgba(0,201,167,0.38) !important;
}
[data-testid="stButton"] > button:active { transform: translateY(0) !important; }

/* ── Download buttons ────────────────────────────────── */
[data-testid="stDownloadButton"] > button {
  background: transparent !important;
  color: var(--teal) !important;
  font-family: 'Syne', sans-serif !important;
  font-weight: 600 !important;
  font-size: 0.9rem !important;
  border: 1.5px solid var(--teal) !important;
  border-radius: 8px !important;
  padding: 0.55rem 1.4rem !important;
  transition: all 0.2s !important;
}
[data-testid="stDownloadButton"] > button:hover {
  background: var(--teal) !important;
  color: var(--navy) !important;
}

/* ── Radio ───────────────────────────────────────────── */
[data-testid="stRadio"] > div {
  gap: 0.6rem !important;
}

/* ── Success / info alerts ───────────────────────────── */
[data-testid="stAlert"] {
  border-radius: var(--radius) !important;
  border-left-width: 4px !important;
}

/* ── Expander ────────────────────────────────────────── */
[data-testid="stExpander"] {
  border: 1.5px solid rgba(138,175,200,0.35) !important;
  border-radius: var(--radius) !important;
  background: rgba(255,255,255,0.04) !important;
}

[data-testid="stExpander"] summary,
[data-testid="stExpander"] summary *,
[data-testid="stExpander"] p,
[data-testid="stExpander"] div {
  color: #dbeafe !important;
}

[data-testid="stExpander"] code,
[data-testid="stExpander"] pre {
  color: #1a2535 !important;
  background: #f8fafc !important;
}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  HERO HEADER
# ─────────────────────────────────────────────

st.markdown("""
<div style="
  background: linear-gradient(135deg, #0d1b2a 0%, #162032 60%, #1a3347 100%);
  border-radius: 16px;
  padding: 36px 40px 32px;
  margin-bottom: 28px;
  position: relative;
  overflow: hidden;
">
  <!-- Decorative grid dots -->
  <div style="
    position:absolute; top:0; right:0; bottom:0; width:340px;
    background-image: radial-gradient(circle, rgba(0,201,167,0.18) 1px, transparent 1px);
    background-size: 22px 22px;
    mask-image: linear-gradient(to left, rgba(0,0,0,0.6), transparent);
    -webkit-mask-image: linear-gradient(to left, rgba(0,0,0,0.6), transparent);
  "></div>
  <!-- Accent blob -->
  <div style="
    position:absolute; top:-60px; right:60px;
    width:220px; height:220px; border-radius:50%;
    background: radial-gradient(circle, rgba(0,201,167,0.10), transparent 70%);
  "></div>

  <div style="display:flex; align-items:center; gap:14px; margin-bottom:10px;">
    <div style="
      background: rgba(0,201,167,0.15);
      border: 1px solid rgba(0,201,167,0.35);
      border-radius: 10px;
      padding: 8px 12px;
      font-size: 1.5rem;
      line-height:1;
    ">🏗️</div>
    <div>
      <div style="
        font-family:'Syne',sans-serif;
        font-size: 1.75rem;
        font-weight: 800;
        color: #ffffff;
        letter-spacing: -0.03em;
        line-height: 1.1;
      ">Kommunal Byggesaksstøtte</div>
      <div style="
        font-size:0.78rem; font-weight:600; letter-spacing:0.12em;
        text-transform:uppercase; color: #00c9a7; margin-top:3px;
      ">Prototype · Plan- og bygningsloven</div>
    </div>
  </div>
  <p style="
    color: #8aafc8;
    font-size: 0.93rem;
    margin: 0;
    max-width: 560px;
    line-height: 1.6;
  ">
    Beslutningsstøtte for enkle byggesaker. Systemet vurderer søknader mot et forenklet regelgrunnlag
    og gir forslag til vedtak — endelig avgjørelse fattes alltid av saksbehandler.
  </p>
</div>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────

with st.sidebar:
    st.markdown("""
    <div style="padding: 8px 0 20px;">
      <div style="font-family:'Syne',sans-serif; font-size:1.1rem; font-weight:700; color:#fff; margin-bottom:4px;">
        Om prototypen
      </div>
      <div style="width:32px; height:3px; background:#00c9a7; border-radius:2px; margin-bottom:16px;"></div>
    </div>
    """, unsafe_allow_html=True)

    st.write("Demoen simulerer kommunal vurdering av enkle byggesaker som garasje og tilbygg.")

    st.markdown("""
    <div style="margin: 16px 0 8px; font-family:'Syne',sans-serif; font-size:0.85rem; font-weight:700; color:#fff; text-transform:uppercase; letter-spacing:0.1em;">
      Formål
    </div>
    """, unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:0.87rem; color:#8aafc8; line-height:1.7;">
      · Beslutningsstøtte for saksbehandlere<br>
      · Tydelig og etterprøvbar regelkontroll<br>
      · Forklarbar, transparent vurdering
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="margin: 20px 0 8px; font-family:'Syne',sans-serif; font-size:0.85rem; font-weight:700; color:#fff; text-transform:uppercase; letter-spacing:0.1em;">
      Avgrensning
    </div>
    """, unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:0.87rem; color:#8aafc8; line-height:1.7;">
      Forenklet regelgrunnlag brukt i bachelorprosjektet. Ikke juridisk bindende.
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="margin-top:32px; padding: 14px; background:rgba(0,201,167,0.08);
                border: 1px solid rgba(0,201,167,0.2); border-radius:10px;">
      <div style="font-size:0.75rem; color:#00c9a7; font-weight:600; text-transform:uppercase; letter-spacing:0.1em; margin-bottom:6px;">
        Versjon
      </div>
      <div style="font-size:0.85rem; color:#8aafc8;">Prototype v1.0<br>Bachelorprosjekt 2025</div>
    </div>
    """, unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  SECTION: REGISTRER SØKNAD
# ─────────────────────────────────────────────

st.markdown("""
<div style="display:flex; align-items:center; gap:10px; margin-bottom:20px;">
  <div style="width:4px; height:28px; background:linear-gradient(180deg,#00c9a7,#0d7c6a); border-radius:2px;"></div>
  <div>
    <div style="font-family:'Syne',sans-serif; font-size:1.2rem; font-weight:700; color:#f8fafc;">
      Registrer byggesøknad
    </div>
    <div style="font-size:0.83rem; color:#8aafc8; margin-top:1px;">
      Fyll inn opplysninger om tiltaket, eller last opp et dokument
    </div>
  </div>
</div>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  DEMO SELECTOR
# ─────────────────────────────────────────────

demo_cases = {
    "Ingen demo valgt": None,
    "✅  Godkjent – garasje innenfor krav": {
        "byggtype": "Garasje", "areal": 40.0, "avstand": 2.0,
        "regulert": "Ja", "dispensasjon": "Nei", "eiendomstype": "Enebolig",
        "mønehøyde": 4.5, "vinduer_mot_nabo": "Nei", "frittstående": "Ja", "nabovarsel_sendt": "Ja",
    },
    "🟡  Krever dispensasjon – for stor garasje": {
        "byggtype": "Garasje", "areal": 68.0, "avstand": 4.0,
        "regulert": "Ja", "dispensasjon": "Ja", "eiendomstype": "Enebolig",
        "mønehøyde": 4.8, "vinduer_mot_nabo": "Nei", "frittstående": "Ja", "nabovarsel_sendt": "Ja",
    },
    "🔴  Avslag – flere avvik": {
        "byggtype": "Garasje", "areal": 70.0, "avstand": 0.5,
        "regulert": "Ja", "dispensasjon": "Nei", "eiendomstype": "Enebolig",
        "mønehøyde": 6.0, "vinduer_mot_nabo": "Ja", "frittstående": "Nei", "nabovarsel_sendt": "Nei",
    },
    "ℹ️  Krever mer info – manglende opplysninger": {
        "byggtype": "Tilbygg", "areal": 0.0, "avstand": 0.0,
        "regulert": "Ja", "dispensasjon": "Nei", "eiendomstype": "Tomannsbolig",
        "mønehøyde": 0.0, "vinduer_mot_nabo": "Nei", "frittstående": "Ja", "nabovarsel_sendt": "Nei",
    },
}

demo_choice = st.selectbox("Velg demo-sak (valgfritt)", list(demo_cases.keys()))
selected_demo = demo_cases.get(demo_choice)

st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

input_mode = st.radio(
    "Inndatamodus",
    ["Fyll inn skjema", "Last opp dokument"],
    horizontal=True,
    index=None,
)

uploaded_data = None
prefill_data = selected_demo.copy() if selected_demo else None

if input_mode is None:
    st.markdown("""
    <div style="
      background: linear-gradient(135deg, #f0f9f7, #e8f4ef);
      border: 1.5px dashed #a7d7cc;
      border-radius: var(--radius);
      padding: 28px;
      text-align: center;
      color: #3d8c7a;
      font-size: 0.95rem;
      margin-top: 12px;
    ">
      <div style="font-size:2rem; margin-bottom:8px;">↑</div>
      Velg inndatamodus over for å starte
    </div>
    """, unsafe_allow_html=True)
    st.stop()

if input_mode == "Last opp dokument":
    st.info("📎 Last opp en .txt-, .json- eller .docx-fil.")
    uploaded_file = st.file_uploader("Last opp byggesøknad", type=["txt", "json", "docx"])

    with st.expander("Se eksempel på filformat"):
        st.code("""byggtype: Garasje\nareal: 45\navstand: 3\nregulert: Ja\ndispensasjon: Ja\neiendomstype: Enebolig\nmønehøyde: 4.5\nvinduer mot nabo: Nei\nfrittstående: Ja\nnabovarsel sendt: Ja""")

    if uploaded_file is not None:
        try:
            uploaded_data = parse_uploaded_application(uploaded_file)
            prefill_data = uploaded_data.copy() if uploaded_data else prefill_data
            st.success("✓ Dokument lest inn. Verdiene under brukes i analysen.")
            with st.expander("Tolket innhold"):
                for k, v in uploaded_data.items():
                    st.write(f"**{k}:** {v}")
        except Exception as e:
            st.error(f"Kunne ikke lese dokumentet: {e}")


# ─────────────────────────────────────────────
#  FORM GRID
# ─────────────────────────────────────────────



col_a, col_b = st.columns(2, gap="large")

with col_a:
    st.markdown("<div style='font-family:Syne,sans-serif; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:0.12em; color:#00a88d; margin-bottom:14px;'>Tiltakets egenskaper</div>", unsafe_allow_html=True)
    byggtype = st.selectbox("Byggtype", ["Garasje", "Tilbygg"],
        index=None if not prefill_data else (["Garasje","Tilbygg"].index(prefill_data.get("byggtype","Garasje")) if prefill_data.get("byggtype","Garasje") in ["Garasje","Tilbygg"] else None),
        placeholder="Velg byggtype")
    areal = st.number_input("Størrelse (m²)", min_value=0.0, max_value=200.0,
        value=float(prefill_data.get("areal", 0)) if prefill_data else None,
        placeholder="Oppgi areal")
    avstand = st.number_input("Avstand til nabogrense (meter)", min_value=0.0, max_value=20.0,
        value=float(prefill_data.get("avstand", 0)) if prefill_data else None,
        placeholder="Oppgi avstand")
    mønehøyde = st.number_input("Mønehøyde (meter)", min_value=0.0, max_value=15.0,
        value=float(prefill_data.get("mønehøyde", 0)) if prefill_data else None,
        placeholder="Oppgi mønehøyde")

with col_b:
    st.markdown("<div style='font-family:Syne,sans-serif; font-size:0.78rem; font-weight:700; text-transform:uppercase; letter-spacing:0.12em; color:#00a88d; margin-bottom:14px;'>Regulering og status</div>", unsafe_allow_html=True)
    regulert = st.selectbox("Er eiendommen regulert?", ["Ja", "Nei"],
        index=None if not prefill_data else (["Ja","Nei"].index(prefill_data.get("regulert","Ja")) if prefill_data.get("regulert","Ja") in ["Ja","Nei"] else None),
        placeholder="Velg status")
    dispensasjon = st.selectbox("Søkes det om dispensasjon?", ["Nei", "Ja"],
        index=None if not prefill_data else (["Nei","Ja"].index(prefill_data.get("dispensasjon","Nei")) if prefill_data.get("dispensasjon","Nei") in ["Nei","Ja"] else None),
        placeholder="Velg ja/nei")
    eiendomstype = st.selectbox("Eiendomstype", ["Enebolig","Tomannsbolig","Fritidsbolig"],
        index=None if not prefill_data else (["Enebolig","Tomannsbolig","Fritidsbolig"].index(prefill_data.get("eiendomstype","Enebolig")) if prefill_data.get("eiendomstype","Enebolig") in ["Enebolig","Tomannsbolig","Fritidsbolig"] else None),
        placeholder="Velg eiendomstype")
    vinduer_mot_nabo = st.selectbox("Vinduer mot nabo?", ["Nei","Ja"],
        index=None if not prefill_data else (["Nei","Ja"].index(prefill_data.get("vinduer_mot_nabo","Nei")) if prefill_data.get("vinduer_mot_nabo","Nei") in ["Nei","Ja"] else None),
        placeholder="Velg ja/nei")
    frittstående = st.selectbox("Er tiltaket frittstående?", ["Ja","Nei"],
        index=None if not prefill_data else (["Ja","Nei"].index(prefill_data.get("frittstående","Ja")) if prefill_data.get("frittstående","Ja") in ["Ja","Nei"] else None),
        placeholder="Velg ja/nei")
    nabovarsel_sendt = st.selectbox("Er nabovarsel sendt?", ["Ja","Nei"],
        index=None if not prefill_data else (["Ja","Nei"].index(prefill_data.get("nabovarsel_sendt","Ja")) if prefill_data.get("nabovarsel_sendt","Ja") in ["Ja","Nei"] else None),
        placeholder="Velg ja/nei")

st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  ANALYSE BUTTON
# ─────────────────────────────────────────────

col_btn, _ = st.columns([1, 3])
with col_btn:
    analyse_btn = st.button("⟶  Analyser søknad", use_container_width=True)

st.divider()


# ─────────────────────────────────────────────
#  RESULTS
# ─────────────────────────────────────────────

if analyse_btn:
    start = time.time()

    data = {
        "byggtype": byggtype, "areal": areal, "avstand": avstand,
        "regulert": regulert, "dispensasjon": dispensasjon, "eiendomstype": eiendomstype,
        "mønehøyde": mønehøyde, "vinduer_mot_nabo": vinduer_mot_nabo,
        "frittstående": frittstående, "nabovarsel_sendt": nabovarsel_sendt,
    }
    result = evaluate_application(data)
    processing_time = round(time.time() - start, 3)
    summary_docx_file = build_summary_docx(data, result)

    # ── Status banner ──────────────────────────────────────
    status_value = str(result.get("indicators", {}).get("status", "")).strip().lower()

    if status_value == "godkjent":
        banner_bg = "linear-gradient(135deg, #064e3b, #065f46)"
        banner_border = "#10b981"
        banner_icon = "✓"
        banner_label = "Godkjent i forenklet vurdering"
        banner_sub = "Søknaden oppfyller alle forenklede hovedkrav"
    elif "dispensasjon" in status_value:
        banner_bg = "linear-gradient(135deg, #78350f, #92400e)"
        banner_border = "#f59e0b"
        banner_icon = "⚠"
        banner_label = "Krever dispensasjon"
        banner_sub = "Søknaden har avvik som krever nærmere vurdering"
    elif "mer info" in status_value:
        banner_bg = "linear-gradient(135deg, #1e3a5f, #1d4ed8)"
        banner_border = "#3b82f6"
        banner_icon = "ℹ"
        banner_label = "Krever mer informasjon"
        banner_sub = "Manglende opplysninger — søknaden må suppleres"
    else:
        banner_bg = "linear-gradient(135deg, #7f1d1d, #991b1b)"
        banner_border = "#ef4444"
        banner_icon = "✕"
        banner_label = "Avslag i forenklet vurdering"
        banner_sub = "Søknaden bryter ett eller flere hovedkrav"

    st.markdown(f"""
    <div style="
      background: {banner_bg};
      border: 1.5px solid {banner_border};
      border-radius: 14px;
      padding: 22px 28px;
      display: flex;
      align-items: center;
      gap: 20px;
      margin-bottom: 24px;
      box-shadow: 0 4px 24px rgba(0,0,0,0.18);
    ">
      <div style="
        width: 52px; height: 52px;
        background: rgba(255,255,255,0.12);
        border: 1.5px solid rgba(255,255,255,0.25);
        border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        font-size: 1.4rem; color: white; flex-shrink: 0;
      ">{banner_icon}</div>
      <div>
        <div style="font-family:'Syne',sans-serif; font-size:1.25rem; font-weight:800; color:white; letter-spacing:-0.02em;">
          {banner_label}
        </div>
        <div style="font-size:0.87rem; color:rgba(255,255,255,0.7); margin-top:3px;">
          {banner_sub}
        </div>
      </div>
      <div style="margin-left:auto; font-family:'Syne',sans-serif; font-size:0.78rem; color:rgba(255,255,255,0.5); text-align:right;">
        Analysert på<br>
        <span style="font-size:1.1rem; color:rgba(255,255,255,0.85); font-weight:700;">{processing_time}s</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Three-column results ───────────────────────────────
    col1, col2, col3 = st.columns(3, gap="small")

    # Card helper
    def card_start(title, icon=""):
        st.markdown(f"""
        <div style="
          background: transparent;
          border: 1.5px solid rgba(138,175,200,0.35);
          border-radius: 14px;
          padding: 14px;
          height: 100%;
          box-shadow: 0 2px 12px rgba(13,27,42,0.06);
          color: #f8fafc;
        ">
          <div style="display:flex; align-items:center; justify-content:center; gap:8px; margin-bottom:12px; padding-bottom:4px; text-align:center; color:#f8fafc;">
            <span style="font-size:1.1rem; color:#f8fafc;">{icon}</span>
            <span style="font-family:'Syne',sans-serif; font-size:0.95rem; font-weight:700; color:#f8fafc;">{title}</span>
          </div>
        """, unsafe_allow_html=True)

    def card_end():
        st.markdown("</div>", unsafe_allow_html=True)

    with col1:
        card_start("Sammendrag av søknad", "📋")
        st.markdown(f"""
        <div style="font-size:0.88rem; color:#dbeafe; line-height:1.7; margin:0;">
          {result["summary"]}
        </div>
        """, unsafe_allow_html=True)
        card_end()

    with col2:
        card_start("Forslag til vedtak", "⚖️")
        # Format the decision text nicely
        decision_lines = result["decision"].split("\n")
        formatted = ""
        for line in decision_lines:
            line = line.strip()
            if not line:
                formatted += "<div style='height:8px'></div>"
            elif line.startswith("Forslag:"):
                formatted += f"<div style='font-family:Syne,sans-serif; font-weight:700; font-size:0.95rem; color:#f8fafc; margin-bottom:6px;'>{line}</div>"
            elif line.startswith("Begrunnelse:") or line.startswith("Løsning") or line.startswith("Anbefaling") or line.startswith("Mulige"):
                formatted += f"<div style='font-size:0.8rem; font-weight:700; text-transform:uppercase; letter-spacing:0.08em; color:#00a88d; margin-top:10px; margin-bottom:4px;'>{line}</div>"
            elif line.startswith("- "):
                formatted += f"<div style='font-size:0.875rem; color:#dbeafe; padding-left:12px; margin-bottom:3px; line-height:1.5;'>· {line[2:]}</div>"
            else:
                formatted += f"<div style='font-size:0.875rem; color:#dbeafe; line-height:1.5;'>{line}</div>"
        st.markdown(f"<div style='color:#dbeafe;'>{formatted}</div>", unsafe_allow_html=True)

        if nabovarsel_sendt == "Nei":
            st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)
            st.markdown("""<div style='font-size:0.82rem; font-weight:600; color:#92400e; background:#fffbeb; border:1px solid #fde68a; border-radius:8px; padding:10px 12px; margin-bottom:8px;'>
              ⚠️ Nabovarsel ikke sendt — last ned mal under
            </div>""", unsafe_allow_html=True)
            nabovarsel_file = build_nabovarsel_template(data)
            st.download_button("📄 Last ned nabovarsel-mal", data=nabovarsel_file,
                file_name="nabovarsel_mal.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        card_end()

    with col3:
        card_start("Regelkontroll & Indikatorer", "🔍")
        indicators = result["indicators"]

        # Rule checks - styled
        st.markdown("<div style='margin-bottom:16px;'>", unsafe_allow_html=True)
        for rule in result.get("rule_checks", []):
            ok = rule.get("ok")
            dot_color = "#10b981" if ok else "#ef4444"
            bg_color = "#f0fdf4" if ok else "#fef2f2"
            border_color = "#bbf7d0" if ok else "#fecaca"
            symbol = "✓" if ok else "✕"
            st.markdown(f"""
            <div style="
              display: flex; align-items: center; gap: 10px;
              background: {bg_color};
              border: 1px solid {border_color};
              border-radius: 8px;
              padding: 8px 12px;
              margin-bottom: 6px;
            ">
              <div style="
                width:20px; height:20px; border-radius:50%;
                background:{dot_color}; color:white;
                display:flex; align-items:center; justify-content:center;
                font-size:0.7rem; font-weight:700; flex-shrink:0;
              ">{symbol}</div>
              <div style="font-size:0.82rem; color:#1f2937; line-height:1.4;">{rule.get('label')}</div>
            </div>
            """, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Indicators
        if indicators.get("antall_avvik"):
            st.markdown(f"""
            <div style="background:#fff7ed; border:1px solid #fed7aa; border-radius:8px; padding:10px 12px; margin-bottom:10px;">
              <div style="font-size:0.75rem; font-weight:700; text-transform:uppercase; letter-spacing:0.08em; color:#c2410c; margin-bottom:3px;">Avvik</div>
              <div style="font-size:0.85rem; color:#7c2d12;">{indicators.get('antall_avvik')}</div>
            </div>
            """, unsafe_allow_html=True)

        for label, key in [("Regelreferanse", "regelreferanse"), ("Konsistens", "konsistens")]:
            val = indicators.get(key, "")
            if val:
                st.markdown(f"""
                <div style="margin-bottom:10px;">
                  <div style="font-size:0.75rem; font-weight:700; text-transform:uppercase; letter-spacing:0.08em; color:#8aafc8; margin-bottom:3px;">{label}</div>
                  <div style="font-size:0.82rem; color:#dbeafe; line-height:1.5;">{val}</div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("""
        <div style="font-size:0.75rem; color:#8aafc8; margin-top:8px; font-style:italic;">
          Lovhenvisninger er forenklede og kun for beslutningsstøtte.
        </div>
        """, unsafe_allow_html=True)
        card_end()

    # ── Download + disclaimer ──────────────────────────────
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    dl_col, disc_col = st.columns([1, 2], gap="large")

    with dl_col:
        st.markdown("""
        <div style="font-family:'Syne',sans-serif; font-size:0.78rem; font-weight:700; text-transform:uppercase;
                    letter-spacing:0.12em; color:#8aafc8; margin-bottom:10px;">
          Last ned resultat
        </div>
        """, unsafe_allow_html=True)
        st.download_button(
            "↓  Last ned som .docx",
            data=summary_docx_file,
            file_name="resultat_byggesak.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with disc_col:
        st.markdown("""
        <div style="
          background: linear-gradient(135deg, #f8fafc, #f1f5f9);
          border: 1.5px solid #e2e8f0;
          border-left: 4px solid #94a3b8;
          border-radius: 10px;
          padding: 14px 18px;
          font-size: 0.83rem;
          color: #475569;
          line-height: 1.7;
        ">
          <strong style="color:#1e293b; font-family:'Syne',sans-serif; font-size:0.82rem; text-transform:uppercase; letter-spacing:0.08em;">
            ⚠ Ansvarsbegrensning
          </strong><br>
          Dette systemet er kun beslutningsstøtte. Endelig vedtak fattes alltid av saksbehandler.
          Resultatet bygger på forenklet regelgrunnlag og <strong>må kontrolleres</strong> mot gjeldende plan,
          lovverk og kommunal praksis.
        </div>
        """, unsafe_allow_html=True)
